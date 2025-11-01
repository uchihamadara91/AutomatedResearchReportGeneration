# research_and_analyst/workflows/report_generator_workflow.py
import os
import sys
import re
from datetime import datetime
from pathlib import Path
import hashlib
from typing import Optional
from jinja2 import Template

from langgraph.graph import StateGraph, START, END
from langgraph.checkpoint.memory import MemorySaver
from langgraph.types import Send

from langchain_core.messages import HumanMessage, SystemMessage

# --- Tavily: prefer new package, fallback to community for compatibility ---
try:
    # pip install -U langchain-tavily
    from langchain_tavily import TavilySearch  # not used here; included for clarity
    from langchain_tavily import TavilySearchResults  # new home
except Exception:
    # fallback if user hasn't upgraded yet
    from langchain_community.tools.tavily_search import TavilySearchResults

from langchain_community.utilities import GoogleSerperAPIWrapper

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

from research_and_analyst.schemas.models import (
    Perspectives,
    GenerateAnalystsState,
    ResearchGraphState,
)
from research_and_analyst.utils.model_loader import ModelLoader
from research_and_analyst.workflows.interview_workflow import InterviewGraphBuilder
from research_and_analyst.prompt_lib.prompt_locator import (
    CREATE_ANALYSTS_PROMPT,
    INTRO_CONCLUSION_INSTRUCTIONS,
    REPORT_WRITER_INSTRUCTIONS,
)
from research_and_analyst.logger import GLOBAL_LOGGER
from research_and_analyst.exception.custom_exception import ResearchAnalystException
from dotenv import load_dotenv

load_dotenv()


class AutonomousReportGenerator:
    """
    End-to-end autonomous report generation workflow using LangGraph.
    - Creates analyst personas
    - Runs multi-source interview pipelines (Tavily + Serper + (optionally arXiv inside InterviewGraph))
    - Consolidates sections into a report (intro + content + conclusion)
    - Saves DOCX/PDF with Windows-safe paths
    """

    def __init__(self, llm):
        self.llm = llm
        self.memory = MemorySaver()
        # Tools: Tavily + Serper (keys are picked from env by these libs)
        self.tavily_search = TavilySearchResults()
        self.serper_search = GoogleSerperAPIWrapper()
        self.logger = GLOBAL_LOGGER.bind(module="AutonomousReportGenerator")

    # ----------------------------------------------------------------------
    # Persona creation
    # ----------------------------------------------------------------------
    def create_analyst(self, state: GenerateAnalystsState):
        """Generate analyst personas based on topic and optional human feedback."""
        topic = state["topic"]
        max_analysts = state["max_analysts"]
        human_analyst_feedback = state.get("human_analyst_feedback", "")

        try:
            self.logger.info("Creating analyst personas", topic=topic)
            structured_llm = self.llm.with_structured_output(Perspectives)
            system_prompt = CREATE_ANALYSTS_PROMPT.render(
                topic=topic,
                max_analysts=max_analysts,
                human_analyst_feedback=human_analyst_feedback,
            )
            analysts = structured_llm.invoke(
                [
                    SystemMessage(content=system_prompt),
                    HumanMessage(content="Generate the set of analysts."),
                ]
            )
            self.logger.info("Analysts created", count=len(analysts.analysts))
            return {"analysts": analysts.analysts}
        except Exception as e:
            self.logger.error("Error creating analysts", error=str(e))
            raise ResearchAnalystException("Failed to create analysts", e)

    # ----------------------------------------------------------------------
    # Human feedback pause
    # ----------------------------------------------------------------------
    def human_feedback(self):
        """Pause node for human analyst feedback (interrupt_before used at compile)."""
        try:
            self.logger.info("Awaiting human feedback")
        except Exception as e:
            self.logger.error("Error during feedback stage", error=str(e))
            raise ResearchAnalystException("Human feedback node failed", e)

    # ----------------------------------------------------------------------
    # Report writing nodes
    # ----------------------------------------------------------------------
    def write_report(self, state: ResearchGraphState):
        """Compile all report sections into unified main content."""
        sections = state.get("sections", [])
        topic = state.get("topic", "")

        try:
            if not sections:
                sections = ["No sections generated — please verify interview stage."]
            self.logger.info("Writing report", topic=topic)
            system_prompt = REPORT_WRITER_INSTRUCTIONS.render(topic=topic)
            report = self.llm.invoke(
                [SystemMessage(content=system_prompt), HumanMessage(content="\n\n".join(sections))]
            )
            self.logger.info("Report written successfully")
            return {"content": report.content}
        except Exception as e:
            self.logger.error("Error writing main report", error=str(e))
            raise ResearchAnalystException("Failed to write main report", e)

    def write_introduction(self, state: ResearchGraphState):
        """Generate the report introduction."""
        try:
            sections = state["sections"]
            topic = state["topic"]
            formatted_str_sections = "\n\n".join([f"{s}" for s in sections])
            self.logger.info("Generating introduction", topic=topic)
            system_prompt = INTRO_CONCLUSION_INSTRUCTIONS.render(
                topic=topic, formatted_str_sections=formatted_str_sections
            )
            intro = self.llm.invoke(
                [SystemMessage(content=system_prompt), HumanMessage(content="Write the report introduction")]
            )
            self.logger.info("Introduction generated", length=len(intro.content))
            return {"introduction": intro.content}
        except Exception as e:
            self.logger.error("Error generating introduction", error=str(e))
            raise ResearchAnalystException("Failed to generate introduction", e)

    def write_conclusion(self, state: ResearchGraphState):
        """Generate the report conclusion."""
        try:
            sections = state["sections"]
            topic = state["topic"]
            formatted_str_sections = "\n\n".join([f"{s}" for s in sections])
            self.logger.info("Generating conclusion", topic=topic)
            system_prompt = INTRO_CONCLUSION_INSTRUCTIONS.render(
                topic=topic, formatted_str_sections=formatted_str_sections
            )
            conclusion = self.llm.invoke(
                [SystemMessage(content=system_prompt), HumanMessage(content="Write the report conclusion")]
            )
            self.logger.info("Conclusion generated", length=len(conclusion.content))
            return {"conclusion": conclusion.content}
        except Exception as e:
            self.logger.error("Error generating conclusion", error=str(e))
            raise ResearchAnalystException("Failed to generate conclusion", e)

    def finalize_report(self, state: ResearchGraphState):
        """Assemble introduction, content, and conclusion into final report."""
        try:
            content = state["content"]
            self.logger.info("Finalizing report compilation")
            if content.startswith("## Insights"):
                # strip heading if present
                content = content.strip("## Insights")

            sources = None
            if "## Sources" in content:
                try:
                    content, sources = content.split("\n## Sources\n")
                except Exception:
                    pass

            final_report = (
                state["introduction"]
                + "\n\n---\n\n"
                + content
                + "\n\n---\n\n"
                + state["conclusion"]
            )
            if sources:
                final_report += "\n\n## Sources\n" + sources

            self.logger.info("Report finalized")
            return {"final_report": final_report}
        except Exception as e:
            self.logger.error("Error finalizing report", error=str(e))
            raise ResearchAnalystException("Failed to finalize report", e)

    # ----------------------------------------------------------------------
    # Windows-safe path helpers (avoid MAX_PATH issues)
    # ----------------------------------------------------------------------
    def _slugify_topic_for_fs(self, topic: str, max_len: int = 40) -> str:
        slug = re.sub(r"[^A-Za-z0-9_\-]+", "_", topic.strip())
        slug = re.sub(r"_+", "_", slug).strip("_")
        return slug[:max_len]

    def _project_root(self) -> str:
        # .../workflows/report_generator_workflow.py -> up two levels
        return os.path.abspath(os.path.join(os.path.dirname(__file__), "../../"))

    def _ensure_dir(self, p: str) -> None:
        Path(p).mkdir(parents=True, exist_ok=True)

    def _build_safe_paths(self, topic: str, fmt: str) -> tuple[str, str, str]:
        """
        Return (root_dir, report_folder, file_path) with path length <= ~240 chars (Windows safe).
        Uses a short filename (e.g., report.docx/pdf) to keep the download name small,
        while keeping uniqueness in the parent folder name.
        """
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        short_hash = hashlib.sha1(topic.encode("utf-8")).hexdigest()[:8]
        slug = self._slugify_topic_for_fs(topic, max_len=40)

        root_dir = os.path.join(self._project_root(), "generated_report")

        base_name = f"{slug}_{short_hash}_{ts}"
        report_folder = os.path.join(root_dir, base_name)
        file_name = f"{base_name}.{fmt}"
        file_path = os.path.join(report_folder, file_name)

        MAX_TOTAL = 240  # under 260 to be safe across libs

        if len(file_path) > MAX_TOTAL:
            reserved = len(os.path.join(root_dir, "")) + len(f"_{short_hash}_{ts}.{fmt}") + 10
            allow_for_slug = max(8, MAX_TOTAL - reserved)
            slug = slug[:allow_for_slug]
            base_name = f"{slug}_{short_hash}_{ts}"
            report_folder = os.path.join(root_dir, base_name)
            file_name = f"{base_name}.{fmt}"
            file_path = os.path.join(report_folder, file_name)

        if len(file_path) > MAX_TOTAL:
            base_name = f"{short_hash}_{ts}"
            report_folder = os.path.join(root_dir, base_name)
            file_name = f"{base_name}.{fmt}"
            file_path = os.path.join(report_folder, file_name)

        # Keep the actual file name short for downloads; uniqueness remains in folder name
        file_name = f"report.{fmt}"
        file_path = os.path.join(report_folder, file_name)

        self.logger.info(
            "Resolved report paths",
            root_dir=root_dir,
            report_folder=report_folder,
            file_name=file_name,
            file_path=file_path,
            path_len=len(file_path),
        )
        return root_dir, report_folder, file_path

    # ----------------------------------------------------------------------
    # Saving (DOCX/PDF)
    # ----------------------------------------------------------------------
    def save_report(self, final_report: str, topic: str, format: str = "docx"):
        """Save the report as DOCX or PDF into generated_report/<slug_hash_ts>/."""
        try:
            self.logger.info("Saving report", topic=topic, format=format)

            _, report_folder, file_path = self._build_safe_paths(topic, format)
            self._ensure_dir(report_folder)

            if not os.path.isdir(report_folder):
                raise FileNotFoundError(f"Target folder was not created: {report_folder}")

            if format == "docx":
                self._save_as_docx(final_report, file_path)
            elif format == "pdf":
                self._save_as_pdf(final_report, file_path)
            else:
                raise ValueError("Invalid format. Use 'docx' or 'pdf'.")

            self.logger.info("Report saved successfully", path=file_path)
            return file_path
        except Exception as e:
            self.logger.error("Error saving report", error=str(e))
            raise ResearchAnalystException("Failed to save report file", e)

    def _save_as_docx(self, text: str, file_path: str):
        """Helper: save as DOCX with simple markdown-ish heading mapping."""
        try:
            doc = Document()
            for line in text.split("\n"):
                if line.startswith("# "):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith("### "):
                    doc.add_heading(line[4:], level=3)
                else:
                    doc.add_paragraph(line)
            doc.save(file_path)
        except Exception as e:
            self.logger.error("DOCX save failed", path=file_path, error=str(e))
            raise ResearchAnalystException("Error saving DOCX report", e)

    def _save_as_pdf(self, text: str, file_path: str):
        """Helper: save as PDF with basic centered layout & wrapping."""
        from textwrap import wrap

        try:
            c = canvas.Canvas(file_path, pagesize=letter)
            width, height = letter

            left_margin = 80
            right_margin = 80
            usable_width = width - left_margin - right_margin
            top_margin = 70
            bottom_margin = 60
            y = height - top_margin

            normal_font = "Helvetica"
            bold_font = "Helvetica-Bold"
            line_height = 15

            lines = text.split("\n")
            for raw_line in lines:
                line = raw_line.strip()
                if not line:
                    y -= line_height
                    continue

                if line.startswith("# "):
                    font = bold_font
                    size = 16
                    line = line[2:].strip()
                elif line.startswith("## "):
                    font = bold_font
                    size = 13
                    line = line[3:].strip()
                else:
                    font = normal_font
                    size = 11

                c.setFont(font, size)
                wrapped_lines = wrap(line, width=int(usable_width / (size * 0.55)))

                for wline in wrapped_lines:
                    if y < bottom_margin:
                        c.showPage()
                        c.setFont(font, size)
                        y = height - top_margin

                    text_width = c.stringWidth(wline, font, size)
                    x = (width - text_width) / 2  # center horizontally
                    c.drawString(x, y, wline)
                    y -= line_height

            # (Optional) footer: page number - reportlab doesn't support total count here easily
            c.save()
            self.logger.info("Centered PDF saved successfully", path=file_path)

        except Exception as e:
            self.logger.error("PDF save failed", path=file_path, error=str(e))
            raise ResearchAnalystException("Error saving PDF report", e)

    # ----------------------------------------------------------------------
    # Graph assembly
    # ----------------------------------------------------------------------
    def build_graph(self):
        """Construct the report generation graph."""
        try:
            self.logger.info("Building report generation graph")
            builder = StateGraph(ResearchGraphState)

            # Build interview subgraph with your multi-search variant
            interview_graph = InterviewGraphBuilder(
                self.llm, self.tavily_search, self.serper_search
            ).build()

            def initiate_all_interviews(state: ResearchGraphState):
                topic = state.get("topic", "Untitled Topic")
                analysts = state.get("analysts", [])
                if not analysts:
                    self.logger.warning("No analysts found — skipping interviews")
                    return END
                return [
                    Send(
                        "conduct_interview",
                        {
                            "analyst": analyst,
                            "messages": [HumanMessage(content=f"So, let's discuss about {topic}.")],
                            "max_num_turns": 2,
                            "context": [],
                            "interview": "",
                            "sections": [],
                        },
                    )
                    for analyst in analysts
                ]

            # Nodes
            builder.add_node("create_analyst", self.create_analyst)
            builder.add_node("human_feedback", self.human_feedback)
            builder.add_node("conduct_interview", interview_graph)
            builder.add_node("write_report", self.write_report)
            builder.add_node("write_introduction", self.write_introduction)
            builder.add_node("write_conclusion", self.write_conclusion)
            builder.add_node("finalize_report", self.finalize_report)

            # Edges
            builder.add_edge(START, "create_analyst")
            builder.add_edge("create_analyst", "human_feedback")
            builder.add_conditional_edges(
                "human_feedback", initiate_all_interviews, ["conduct_interview", END]
            )
            builder.add_edge("conduct_interview", "write_report")
            builder.add_edge("conduct_interview", "write_introduction")
            builder.add_edge("conduct_interview", "write_conclusion")
            builder.add_edge(
                ["write_report", "write_introduction", "write_conclusion"],
                "finalize_report",
            )
            builder.add_edge("finalize_report", END)

            graph = builder.compile(
                interrupt_before=["human_feedback"], checkpointer=self.memory
            )
            self.logger.info("Report generation graph built successfully")
            return graph
        except Exception as e:
            self.logger.error("Error building report graph", error=str(e))
            raise ResearchAnalystException("Failed to build report generation graph", e)


# ----------------------------------------------------------------------
# Standalone smoke test (optional)
# ----------------------------------------------------------------------
if __name__ == "__main__":
    try:
        llm = ModelLoader().load_llm()
        reporter = AutonomousReportGenerator(llm)
        graph = reporter.build_graph()

        topic = "Impact of LLMs over the Future of Jobs?"
        thread = {"configurable": {"thread_id": "1"}}
        reporter.logger.info("Starting report generation pipeline", topic=topic)

        for _ in graph.stream({"topic": topic, "max_analysts": 3}, thread, stream_mode="values"):
            pass

        state = graph.get_state(thread)
        feedback = input("\n Enter your feedback or press Enter to continue: ").strip()
        graph.update_state(thread, {"human_analyst_feedback": feedback}, as_node="human_feedback")

        for _ in graph.stream(None, thread, stream_mode="values"):
            pass

        final_state = graph.get_state(thread)
        final_report = final_state.values.get("final_report")

        if final_report:
            reporter.logger.info("Report generated successfully")
            reporter.save_report(final_report, topic, "docx")
            reporter.save_report(final_report, topic, "pdf")
        else:
            reporter.logger.error("No report content generated")

    except Exception as e:
        GLOBAL_LOGGER.error("Fatal error in main execution", error=str(e))
        raise ResearchAnalystException("Autonomous report generation pipeline failed", e)
